from fastapi import UploadFile, HTTPException
from docx import Document
import os
import uuid
from typing import Dict, Any
import json

class DocumentService:
    def __init__(self):
        self.upload_dir = "uploads"
        self.processed_dir = "processed"
        os.makedirs(self.upload_dir, exist_ok=True)
        os.makedirs(self.processed_dir, exist_ok=True)

    async def process_upload(self, file: UploadFile) -> Dict[str, Any]:
        """Process the uploaded document"""
        document_id = str(uuid.uuid4())
        file_path = os.path.join(self.upload_dir, f"{document_id}.docx")
        
        # Save the uploaded file
        with open(file_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)
        
        # Extract text from the document
        doc = Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        # Save metadata
        metadata = {
            "original_filename": file.filename,
            "document_id": document_id,
            "status": "uploaded"
        }
        
        with open(os.path.join(self.upload_dir, f"{document_id}_metadata.json"), "w") as f:
            json.dump(metadata, f)
        
        return {
            "document_id": document_id,
            "message": "Document uploaded successfully",
            "text": text
        }

    async def get_document(self, document_id: str, clean: bool = False) -> Dict[str, Any]:
        """Get the processed document"""
        file_path = os.path.join(
            self.processed_dir if clean else self.upload_dir,
            f"{document_id}.docx"
        )
        
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="Document not found")
        
        return {
            "file_path": file_path,
            "document_id": document_id
        }

    def create_redline_document(self, document_id: str, changes: Dict[str, Any]) -> str:
        """Create a redline version of the document with suggested changes"""
        original_path = os.path.join(self.upload_dir, f"{document_id}.docx")
        output_path = os.path.join(self.processed_dir, f"{document_id}_redline.docx")
        
        doc = Document(original_path)
        
        # Apply changes to the document
        for paragraph in doc.paragraphs:
            if paragraph.text in changes:
                # Apply redline formatting
                paragraph.text = changes[paragraph.text]["suggestion"]
                for run in paragraph.runs:
                    run.font.color.rgb = (255, 0, 0)  # Red color
        
        doc.save(output_path)
        return output_path

    def create_clean_document(self, document_id: str) -> str:
        """Create a clean version of the document with accepted changes"""
        redline_path = os.path.join(self.processed_dir, f"{document_id}_redline.docx")
        output_path = os.path.join(self.processed_dir, f"{document_id}_clean.docx")
        
        doc = Document(redline_path)
        
        # Remove redline formatting and keep the suggested changes
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = (0, 0, 0)  # Black color
        
        doc.save(output_path)
        return output_path 